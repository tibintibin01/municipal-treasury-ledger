import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "USER_MANUAL.md"
OUTPUT = ROOT / "Business_Tax_Permit_Collection_System_User_Manual.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_YELLOW = "FFF8E1"
GRAY = "666666"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B7C4D4")
        borders.append(tag)


def apply_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    if bold:
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        apply_cell_text(cell, header, True)
        if widths:
            cell.width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            apply_cell_text(cells[i], value)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_note_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_YELLOW)
    set_cell_margins(cell, 120, 120, 160, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_screenshot_placeholder(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, 180, 180, 180, 180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(102, 102, 102)
    doc.add_paragraph()


def format_inline(paragraph, text):
    token_re = re.compile(r"(`[^`]+`|\\*\\*[^*]+\\*\\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_docx():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)

    # Cover block.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Business Tax & Permit Collection System")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("User Manual")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Version v0.3.44\\nPrepared for Municipal Treasurer's Office staff")
    doc.add_paragraph()

    add_note_box(
        doc,
        "This manual is written for beginner users. Follow the steps in order and ask an Admin or Treasurer before restoring backups, deleting records, or changing system settings.",
    )
    doc.add_page_break()

    # Skip first four lines already used as cover.
    current_numbered = False
    for raw in lines[6:]:
        line = raw.rstrip()
        if not line:
            current_numbered = False
            continue
        if line.startswith("[Screenshot:"):
            add_screenshot_placeholder(doc, line)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            format_inline(p, line[2:].strip())
            current_numbered = False
            continue
        if re.match(r"^\\d+\\.\\s+", line):
            p = doc.add_paragraph(style="List Number")
            format_inline(p, re.sub(r"^\\d+\\.\\s+", "", line).strip())
            current_numbered = True
            continue
        if line.startswith("Important:") or line.startswith("Note:"):
            add_note_box(doc, line)
            current_numbered = False
            continue
        p = doc.add_paragraph()
        format_inline(p, line)
        current_numbered = False

    # Add compact reference tables at the end for print users.
    doc.add_page_break()
    doc.add_heading("Quick Reference Tables", level=1)
    add_table(
        doc,
        ["Role", "Main Use"],
        [
            ["Admin", "Manages users, settings, backups, audit logs, and records."],
            ["Treasurer", "Manages treasury records, reports, backups, and settings."],
            ["Cashier", "Records owner, assessment, and payment information based on allowed access."],
        ],
        [Inches(1.6), Inches(4.9)],
    )
    add_table(
        doc,
        ["Button", "Purpose"],
        [
            ["Password", "Change the current user's password."],
            ["Theme", "Switch light or dark mode. Reopen the app to apply."],
            ["Save Data", "Manually save ledger data."],
            ["Backup To", "Choose the backup folder."],
            ["Backup", "Create and verify a backup."],
            ["Restore", "Restore from a backup file."],
            ["Import", "Import Excel or CSV records."],
            ["Export", "Export ledger data to Excel or CSV."],
        ],
        [Inches(1.6), Inches(4.9)],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Business Tax & Permit Collection System User Manual")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
